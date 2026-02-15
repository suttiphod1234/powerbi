# ===============================================
# ADVANCED DAX WORKSHOP – COMPLETE PROJECT FILES
# ===============================================

from openpyxl import Workbook
from datetime import datetime, timedelta
import random
from docx import Document
import zipfile
import os

# ------------------------------------------------
# 1) CREATE DATASET (Excel)
# ------------------------------------------------

wb = Workbook()

# Sales Sheet
ws_sales = wb.active
ws_sales.title = "Sales"
ws_sales.append(["OrderID", "Date", "ProductID", "CustomerID", "Quantity", "UnitPrice"])

products = ["P01", "P02", "P03", "P04"]
customers = ["C01", "C02", "C03", "C04", "C05"]
prices = {"P01": 500, "P02": 1200, "P03": 800, "P04": 300}

start_date = datetime(2023, 1, 1)
order_id = 1001

for i in range(1000):
    date = start_date + timedelta(days=random.randint(0, 730))
    product = random.choice(products)
    customer = random.choice(customers)
    qty = random.randint(1, 5)
    ws_sales.append([order_id, date.date(), product, customer, qty, prices[product]])
    order_id += 1

# Product Sheet
ws_product = wb.create_sheet("Product")
ws_product.append(["ProductID", "Product Name", "Category"])
ws_product.append(["P01", "Laptop", "Electronics"])
ws_product.append(["P02", "Printer", "Electronics"])
ws_product.append(["P03", "Desk", "Furniture"])
ws_product.append(["P04", "Chair", "Furniture"])

# Customer Sheet
ws_customer = wb.create_sheet("Customer")
ws_customer.append(["CustomerID", "Customer Name", "Region"])
ws_customer.append(["C01", "ABC Co", "Bangkok"])
ws_customer.append(["C02", "XYZ Ltd", "Chiang Mai"])
ws_customer.append(["C03", "Smart Biz", "Bangkok"])
ws_customer.append(["C04", "Good Tech", "Phuket"])
ws_customer.append(["C05", "Modern Co", "Khon Kaen"])

# Date Table
ws_date = wb.create_sheet("Date")
ws_date.append(["Date", "Year", "Month", "MonthNo", "YearMonth"])

current_date = datetime(2023, 1, 1)
end_date = datetime(2024, 12, 31)

while current_date <= end_date:
    ws_date.append([
        current_date.date(),
        current_date.year,
        current_date.strftime("%B"),
        current_date.month,
        f"{current_date.year}-{current_date.month:02d}"
    ])
    current_date += timedelta(days=1)

dataset_path = "Advanced_DAX_Project_Dataset.xlsx"
wb.save(dataset_path)

# ------------------------------------------------
# 2) CREATE MEASURES FILE (TXT)
# ------------------------------------------------

measures_text = """
==============================
ADVANCED DAX – COMPLETE MEASURES
==============================

Total Sales =
SUMX(Sales, Sales[Quantity] * Sales[UnitPrice])

Total Profit =
[Total Sales] * 0.3

Sales 2024 =
CALCULATE([Total Sales], 'Date'[Year] = 2024)

YTD Sales =
TOTALYTD([Total Sales], 'Date'[Date])

Sales LY =
CALCULATE([Total Sales], SAMEPERIODLASTYEAR('Date'[Date]))

YoY Growth % =
DIVIDE([Total Sales] - [Sales LY], [Sales LY])

Rolling 6M =
CALCULATE(
    [Total Sales],
    DATESINPERIOD(
        'Date'[Date],
        MAX('Date'[Date]),
        -6,
        MONTH
    )
)

% of Total =
DIVIDE(
    [Total Sales],
    CALCULATE([Total Sales], ALL(Product))
)

Ranking =
RANKX(
    ALL(Product[Product Name]),
    [Total Sales]
)
"""

measures_path = "Advanced_DAX_Measures.txt"
with open(measures_path, "w", encoding="utf-8") as f:
    f.write(measures_text)

# ------------------------------------------------
# 3) CREATE FINAL EXAM (DOCX)
# ------------------------------------------------

doc = Document()
doc.add_heading("Advanced DAX Techniques – Final Exam", level=1)

doc.add_heading("Section A: Practical (70 คะแนน)", level=2)
doc.add_paragraph("1) สร้าง Measure: YTD Sales")
doc.add_paragraph("2) สร้าง Measure: YoY Growth %")
doc.add_paragraph("3) สร้าง Rolling 3 Months")
doc.add_paragraph("4) แสดง Top 5 Product พร้อม %Contribution")
doc.add_paragraph("5) สร้าง Ranking ตาม Region")

doc.add_heading("Section B: Concept (30 คะแนน)", level=2)
doc.add_paragraph("6) อธิบายความแตกต่างระหว่าง Row Context และ Filter Context")
doc.add_paragraph("7) CALCULATE ทำงานอย่างไร")
doc.add_paragraph("8) ALL กับ REMOVEFILTERS ต่างกันอย่างไร")

doc.add_page_break()

doc.add_heading("=== ANSWER KEY ===", level=1)

doc.add_paragraph("Rolling 3M =")
doc.add_paragraph("""
CALCULATE(
    [Total Sales],
    DATESINPERIOD(
        'Date'[Date],
        MAX('Date'[Date]),
        -3,
        MONTH
    )
)
""")

doc.add_paragraph("Top 5 Example =")
doc.add_paragraph("""
CALCULATE(
    [Total Sales],
    TOPN(
        5,
        SUMMARIZE(
            Sales,
            Product[Product Name],
            "Total", [Total Sales]
        ),
        [Total],
        DESC
    )
)
""")

exam_path = "Advanced_DAX_Final_Exam.docx"
doc.save(exam_path)

# ------------------------------------------------
# 4) ZIP ALL FILES
# ------------------------------------------------

zip_path = "Advanced_DAX_Complete_Project.zip"

with zipfile.ZipFile(zip_path, 'w') as zipf:
    zipf.write(dataset_path, dataset_path)
    zipf.write(measures_path, measures_path)
    zipf.write(exam_path, exam_path)

print(f"Generated: {dataset_path}, {measures_path}, {exam_path}, {zip_path}")
